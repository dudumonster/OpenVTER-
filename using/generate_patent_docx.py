#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Convert a UTF-8 markdown-like patent draft to docx.

Rules:
- "# " => document title
- "## " => level 1 heading
- other non-empty lines => paragraphs
"""

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def parse_args():
    parser = argparse.ArgumentParser(description="Generate patent docx from markdown-like text.")
    parser.add_argument("--input", required=True, help="Input markdown-like UTF-8 file")
    parser.add_argument("--output", required=True, help="Output .docx file")
    return parser.parse_args()


def set_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(15)


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(13)


def add_paragraph(doc: Document, text: str):
    doc.add_paragraph(text)


def main():
    args = parse_args()
    in_file = Path(args.input)
    out_file = Path(args.output)
    out_file.parent.mkdir(parents=True, exist_ok=True)

    lines = in_file.read_text(encoding="utf-8").splitlines()

    doc = Document()
    set_default_font(doc)

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            add_title(doc, line[2:].strip())
        elif line.startswith("## "):
            add_h1(doc, line[3:].strip())
        else:
            add_paragraph(doc, line)

    doc.save(str(out_file))
    print(str(out_file))


if __name__ == "__main__":
    main()

